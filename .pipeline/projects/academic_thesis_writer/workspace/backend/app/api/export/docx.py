"""DOCX exporter for thesis drafts."""

from __future__ import annotations

import logging
from typing import Optional

from ..models import Draft

logger = logging.getLogger(__name__)


class DocxExporter:
    """Export a thesis draft as a DOCX file."""

    @staticmethod
    def export(draft: Draft) -> bytes:
        """Export the draft as a DOCX bytes object."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

        doc = Document()

        # Title
        doc.add_heading(draft.title, level=0)
        doc.add_paragraph(f"Topic: {draft.topic}")
        doc.add_paragraph(f"Citation Style: {draft.citation_style.value}")

        # Abstract
        if draft.abstract:
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(draft.abstract)

        # Sections
        for section in draft.sections:
            doc.add_heading(section.name, level=1)
            # Split content into paragraphs
            paragraphs = section.content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Bibliography
        if draft.bibliography:
            doc.add_heading("Bibliography", level=1)
            for entry in draft.bibliography:
                doc.add_paragraph(entry.formatted, style="List Bullet")

        # Save to bytes
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def save(draft: Draft, file_path: str) -> None:
        """Save the draft as a DOCX file."""
        content = DocxExporter.export(draft)
        with open(file_path, "wb") as f:
            f.write(content)
        logger.info("Saved DOCX to %s", file_path)
